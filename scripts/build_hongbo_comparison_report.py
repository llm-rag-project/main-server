import argparse
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


FONT = "Malgun Gothic"
ACCENT = "2E74B5"
ACCENT_DARK = "1F4D78"
TEXT = "1F2937"
MUTED = "667085"
LINE = "D0D5DD"
HEADER_FILL = "F2F4F7"
SOFT_BLUE = "EFF6FF"
SOFT_GREEN = "ECFDF3"
SOFT_AMBER = "FFFAEB"
SOFT_RED = "FEF3F2"
WHITE = "FFFFFF"

PORTRAIT_WIDTH_DXA = 9360
LANDSCAPE_WIDTH_DXA = 13680
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--comparison-json", type=Path, required=True)
    parser.add_argument("--baseline-json", type=Path)
    parser.add_argument("--natural-audit-json", type=Path)
    parser.add_argument("--diagnostics-json", type=Path)
    parser.add_argument("--archive-backfill-json", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


def load_json(path: Path | None) -> dict:
    if path is None or not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def set_run_font(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = FONT
    r_pr = run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        r_pr.rFonts.set(qn(f"w:{key}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, color: str) -> None:
    return None


def set_cell_margins(cell, margins=None) -> None:
    return None


def set_table_geometry(table, widths: list[int], total_width: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, width in enumerate(widths):
        table.columns[index].width = Twips(width)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Twips(widths[index])


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def keep_paragraph_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def keep_row_together(row) -> None:
    return None


def style_table(table, widths: list[int], total_width: int, font_size=8.5) -> None:
    table.style = "Light Shading Accent 1"
    set_table_geometry(table, widths, total_width)
    repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        keep_row_together(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        font_size,
                        row_index == 0,
                        ACCENT_DARK if row_index == 0 else TEXT,
                    )


def add_heading(document, text: str, level: int = 1):
    style = document.styles[f"Heading {level}"]
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    return paragraph


def add_body(document, text: str, *, bold_prefix: str | None = None):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, 11, True, TEXT)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(rest, 11, False, TEXT)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 11, False, TEXT)
    return paragraph


def add_bullet(document, text: str):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, TEXT)
    return paragraph


def add_numbered_step(document, title: str, body: str):
    paragraph = document.add_paragraph(style="List Number")
    title_run = paragraph.add_run(title)
    set_run_font(title_run, 11, True, ACCENT_DARK)
    body_run = paragraph.add_run(f"  {body}")
    set_run_font(body_run, 11, False, TEXT)
    return paragraph


def add_callout(document, label: str, text: str, fill=SOFT_BLUE) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, 11, True, ACCENT_DARK)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, 11, False, TEXT)


def add_table(document, headers: list[str], rows: list[list], widths: list[int], *, total_width=PORTRAIT_WIDTH_DXA, font_size=8.5):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.cell(0, index).text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    style_table(table, widths, total_width, font_size=font_size)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("동국대학교 홍보처 비교분석")
    set_run_font(label, 8, False, MUTED)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, ACCENT, 16, 8),
        2: (13, ACCENT, 12, 6),
        3: (12, ACCENT_DARK, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section, *, landscape=False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def percent_change(current: float, baseline: float) -> str:
    return f"{current - baseline:+.1f}%p"


def main() -> None:
    args = parse_args()
    comparison = load_json(args.comparison_json)
    baseline = load_json(args.baseline_json)
    evaluation = comparison["evaluation"]
    metrics = evaluation["metrics"]
    summary = comparison["summary"]
    baseline_metrics = (baseline.get("evaluation") or {}).get("metrics") or {}
    baseline_summary = baseline.get("summary") or {}

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("동국대학교 홍보처 기존 메일")
    set_run_font(run, 23, True, TEXT)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("기사 수집·AI 선정 비교분석 보고서")
    set_run_font(run, 17, True, ACCENT)

    metadata = [
        ("분석 범위", "2025.04.08 ~ 2026.07.10"),
        ("비교 대상", f"기존 메일 {summary['date_count']}일 · 원본 기사 {summary['original_count']}건"),
        ("재검증 기준일", date.today().isoformat()),
        ("평가 원칙", "기존 홍보처 원본을 AI 후보에 강제로 넣지 않은 독립 비교"),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, 10.5, True, TEXT)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, 10.5, False, TEXT)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(9)
    rule.paragraph_format.space_after = Pt(14)
    run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(run, 9, False, ACCENT)

    add_callout(
        document,
        "한눈에 보는 결론",
        (
            f"과거 원본 {metrics['original_count']}건은 모두 DB에 보존됐습니다. "
            f"현재 발송일 규칙 안에서 비교 가능한 후보는 {metrics['candidate_available_count']}건({metrics['candidate_recall_rate']:.1f}%)이며, "
            f"그 후보 중 AI가 같은 URL 또는 같은 주제의 대표 기사를 선택한 비율은 {metrics['candidate_selection_recall_rate']:.1f}%입니다. "
            f"원본 전체 기준 최종 재현율은 {metrics['end_to_end_match_rate']:.1f}%입니다."
        ),
        SOFT_BLUE,
    )

    add_heading(document, "1. 핵심 결과")
    kpi_rows = [[
        f"{metrics['archive_stored_count']}/{metrics['original_count']}\n{metrics['archive_stored_rate']:.1f}%",
        f"{metrics['candidate_available_count']}/{metrics['original_count']}\n{metrics['candidate_recall_rate']:.1f}%",
        f"{metrics['topic_inclusive_match_count']}/{metrics['candidate_available_count']}\n{metrics['candidate_selection_recall_rate']:.1f}%",
        f"{metrics['topic_inclusive_match_count']}/{metrics['selected_count']}\n{metrics['selection_precision_rate']:.1f}%",
        f"{metrics['topic_inclusive_match_count']}/{metrics['original_count']}\n{metrics['end_to_end_match_rate']:.1f}%",
    ]]
    kpi = add_table(
        document,
        ["원본 DB 보존", "날짜 범위 후보 확보", "후보 내 AI 재현", "선택 기사 정확도", "전체 재현율"],
        kpi_rows,
        [1872] * 5,
        font_size=8.2,
    )
    for cell in kpi.rows[1].cells:
        set_cell_fill(cell, SOFT_BLUE)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, 12, True, ACCENT_DARK)

    add_heading(document, "개선 전후", level=2)
    before_end_to_end = float(baseline_metrics.get("topic_match_rate", baseline_summary.get("historical_recall_rate", 0)))
    before_candidate_ai = float(
        baseline_metrics.get(
            "candidate_selection_recall_rate",
            baseline_metrics.get("policy_adjusted_topic_match_rate", 0),
        )
    )
    before_precision = float(baseline_metrics.get("selection_precision_rate", baseline_summary.get("selection_precision_rate", 0)))
    before_matches = int(baseline_metrics.get("topic_inclusive_match_count", baseline_summary.get("total_match_count", 0)))
    improvement_rows = [
        ["최종 일치 기사", f"{before_matches}건", f"{metrics['topic_inclusive_match_count']}건", f"{metrics['topic_inclusive_match_count'] - before_matches:+d}건"],
        ["원본 전체 재현율", f"{before_end_to_end:.1f}%", f"{metrics['end_to_end_match_rate']:.1f}%", percent_change(metrics["end_to_end_match_rate"], before_end_to_end)],
        ["후보 내 AI 재현율", f"{before_candidate_ai:.1f}%", f"{metrics['candidate_selection_recall_rate']:.1f}%", percent_change(metrics["candidate_selection_recall_rate"], before_candidate_ai)],
        ["선택 기사 정확도", f"{before_precision:.1f}%", f"{metrics['selection_precision_rate']:.1f}%", percent_change(metrics["selection_precision_rate"], before_precision)],
    ]
    add_table(document, ["지표", "개선 전", "현재", "변화"], improvement_rows, [3420, 1800, 1800, 2340])
    add_body(
        document,
        "선택 기사 정확도는 1.0%p 낮아졌지만 최종 선정 기사 수가 439건에서 475건으로 늘고, 기존 홍보처와 일치한 기사는 26건 증가했습니다. 즉 더 넓게 후보를 살리면서도 선택 결과의 84.6%가 기존 홍보처 판단과 일치했습니다.",
    )
    add_callout(
        document,
        "주의",
        "이전에 표시된 70.4%는 기존 홍보처 원본을 AI 후보군에 포함했던 검증 결과이므로 실제 독립 선정 성능으로 사용할 수 없습니다. 이번 보고서는 원본을 후보에 강제로 넣지 않은 55.7%를 정직한 개선 전 기준으로 사용합니다.",
        SOFT_AMBER,
    )

    add_heading(document, "2. 서버는 어떤 기사를 수집하는가")
    add_body(
        document,
        "서버는 ‘동국대학교’라는 글자가 제목에 있는 기사만 저장하지 않습니다. 검색 서버가 여러 소스에서 후보를 모은 뒤, 날짜·원문 주소·관련성·중복 여부를 차례로 확인합니다.",
    )
    add_numbered_step(document, "발행 시간 범위를 정합니다.", "기본 메일 발송 시각은 오전 8시 30분입니다. 선택한 업무일의 전날 오전 8시 30분부터 당일 오전 8시 30분까지 발행된 기사를 확인합니다.")
    add_numbered_step(document, "주말과 휴일을 자동으로 연결합니다.", "월요일에는 토·일·월 범위를, 공휴일·학교 휴일·개인 휴가 다음 업무일에는 휴일 기간 전체를 자동으로 포함합니다.")
    add_numbered_step(document, "여러 수집 경로를 함께 사용합니다.", "Naver, Google News RSS, 동국대 공식 소스, 관련 기사 확장, 교육 전용 후보군, 불교 전용 후보군을 하나의 수집 파이프라인으로 합칩니다.")
    add_numbered_step(document, "관련성을 넓게 확인합니다.", "제목과 요약을 먼저 보고, 동국대가 없으면 원문을 보강해 발행인·언론사·기자명·키워드·주석·게시판 정보와 본문을 확인합니다. 본문에만 등장하는 경우에는 오탐을 줄이기 위해 동국대 또는 동국대학교 표현이 두 번 이상 확인돼야 합니다.")
    add_numbered_step(document, "교육·종단 기사는 별도 기준으로 확보합니다.", "대학 [교육]과 불교 [종단] 전용 후보군은 동국대 직접 언급이 없어도 해당 섹션의 정책·종단 동향으로 수집할 수 있습니다.")
    add_numbered_step(document, "실패하면 다시 시도합니다.", "검색 서버의 연결 오류와 시간 초과는 한 작업 안에서 최대 3회 재시도하고, 정기 수집 결과가 불완전하면 자동 재수집 작업도 최대 3회까지 실행합니다.")
    add_numbered_step(document, "같은 원문만 1건으로 정리합니다.", "추적 파라미터를 제거한 정규화 URL과 본문 지문을 사용해 완전히 같은 기사를 통합합니다. 같은 주제를 다룬 서로 다른 언론사 기사는 이 단계에서 지우지 않고 AI 판단 대상으로 남깁니다.")

    add_heading(document, "수집 후 세 개 상위 분류", level=2)
    collection_rows = [
        ["동국대 [법인/건학위]", "동국대학교·법인·건학위원회·교내기관·동문·교수와 직접 연결된 기사"],
        ["대학 [교육]", "교육부, 대교협, 사총협, 고등교육법, 등록금, 입시 등 고등교육 정책과 교육 이슈"],
        ["불교 [종단]", "조계종, 종단, 포교, 출가, 성보, 불교문화유산 등 불교계 주요 동향"],
    ]
    add_table(document, ["상위 분류", "수집 범위"], collection_rows, [2700, 6660])

    add_heading(document, "3. AI는 어떤 기준과 과정으로 기사를 선별하는가")
    add_body(
        document,
        "AI는 ‘점수가 높은 순서대로 자르는 도구’가 아니라, 같은 사건을 묶고 대표 기사 하나를 고른 뒤 홍보처 기준에 맞게 분류·요약·정렬하는 편집자 역할을 합니다.",
    )
    ai_steps = [
        ("후보 전달", "제목, 언론사, 기존 요약, 원문 URL, 정규화 URL, 관련 링크, 작성일, 상위·하위 분류를 AI에 전달합니다."),
        ("제외 기준 적용", "관리자가 지정한 제외 문장에 따라 인사발령, 특정 캠퍼스, 원문 확인 불가 기사처럼 원하지 않는 대상을 먼저 제외합니다."),
        ("동일 주제 묶음", "제목이 달라도 인물·기관·행사·금액·성과·본문 내용이 같은 보도자료나 같은 사건이면 하나의 그룹으로 묶습니다."),
        ("대표 기사 1건 선정", "정상 원문 URL, 본문 충실도, 제목의 명확성, 핵심 사실의 완전성, 언론사 신뢰도, 홍보처 활용도를 비교해 대표 기사 하나를 고릅니다. 나머지 매체 링크는 관련 링크로 보존합니다."),
        ("우선순위 판단", "총장·이사장 메시지, 기부·장학·발전기금, 진행 중 캠페인, 공식 보도자료·행사, 연구·AI, 협약, 수상 등을 우선합니다. 관리자가 문장을 수정하면 수정된 전문을 그대로 가장 우선적인 기준으로 사용합니다."),
        ("메일용 정보 생성", "대표 기사마다 자연스러운 1~2문장 요약, 상위·하위 카테고리, 우선순위, 선정 이유, 대표 URL과 관련 링크를 반환합니다."),
        ("서버 최종 점검", "이전 발송 주제, 완전 중복, 섹션 자격, 카테고리별 최대 기사 수를 확인하고 최종 메일 초안을 DB에 저장합니다. 사용자가 순서·요약·포함 여부를 고치면 그 초안이 계속 유지됩니다."),
    ]
    for title_text, body_text in ai_steps:
        add_numbered_step(document, title_text, body_text)

    add_heading(document, "현재 기본 우선순위", level=2)
    add_bullet(document, "가장 우선: 총장·이사장 메시지, 기부·장학·발전기금, 건학 120주년 등 진행 캠페인, 학교 공식 보도자료·행사")
    add_bullet(document, "주요 성과: 연구·기술 개발·특허·AI, 공식 협약·사업 선정, 수상·인증")
    add_bullet(document, "일반·참고: 학술대회, 입시·교육 프로그램, 인사, 고등교육 정책, 동문·교수 인터뷰, 불교계 일반 소식")
    add_bullet(document, "대표 기사: URL 정상 여부, 본문 충실도, 제목 명확성, 핵심 정보 완전성, 언론사 신뢰도, 홍보 활용도")
    add_body(
        document,
        "운영 중에는 사용자가 설정 화면에서 우선순위 기준, 제외 기준, 중복 주제 대표 기사 선정 기준을 문장으로 수정할 수 있습니다. 월별 사용자 편집 이력은 AI 인사이트로 분석되어 소폭 보정되고, 분기별로 전체 기준을 재조정할 수 있습니다.",
    )

    add_heading(document, "4. 비교 방법과 지표를 분리한 이유")
    methodology_rows = [
        ["원본 DB 보존율", "과거 홍보처 원본이 현재 DB에 남아 있는가", "수집 이력 보존 상태"],
        ["날짜 범위 후보 확보율", "발송일의 업무일 규칙 안에서 원본이 후보로 보이는가", "수집·날짜 규칙 성능"],
        ["후보 내 AI 재현율", "후보로 확보된 기사 중 AI가 같은 URL 또는 주제를 고르는가", "AI 선정 성능"],
        ["선택 기사 정확도", "AI가 실제 선택한 기사 중 기존 홍보처와 같은 URL 또는 주제인가", "선택 결과의 정밀도"],
        ["원본 전체 재현율", "원본 전체 중 최종적으로 같은 URL 또는 주제가 선택됐는가", "수집과 AI를 합친 전체 성능"],
    ]
    add_table(document, ["지표", "질문", "의미"], methodology_rows, [2200, 4380, 2780], font_size=8.2)
    add_body(
        document,
        "URL 일치만으로 비교하면 같은 보도자료를 다른 언론사가 대표로 보도한 정상적인 선택을 실패로 보게 됩니다. 그래서 ‘URL 일치’, ‘정규화 제목 일치’, ‘동일 주제의 다른 대표 기사’를 분리하고, 운영 관점의 최종 결과에는 세 유형을 함께 표시했습니다.",
    )

    add_heading(document, "5. 전체 및 카테고리별 결과")
    category_rows = []
    for row in evaluation["category_metrics"]:
        category_rows.append([
            row["label"],
            row["original_count"],
            f"{row['archive_stored_rate']:.1f}%",
            f"{row['candidate_recall_rate']:.1f}%",
            f"{row['candidate_selection_recall_rate']:.1f}%",
            f"{row['topic_match_rate']:.1f}%",
        ])
    add_table(
        document,
        ["상위 분류", "원본", "DB 보존", "후보 확보", "후보 내 AI 재현", "전체 재현"],
        category_rows,
        [2700, 900, 1260, 1440, 1620, 1440],
        font_size=8.2,
    )
    add_bullet(document, f"동국대 [법인/건학위]: 후보 확보 후 AI 재현율 73.5%. 직접 홍보 기사 종류가 넓고, 행사·동문·교수 기사 간 우선순위 차이가 커 가장 개선 여지가 큽니다.")
    add_bullet(document, f"대학 [교육]: 후보 확보율은 68.5%로 가장 낮지만 후보가 확보되면 AI 재현율은 84.3%입니다. 핵심 과제는 AI보다 교육 후보 수집 확대입니다.")
    add_bullet(document, f"불교 [종단]: 후보 확보율 77.5%, 후보 내 AI 재현율 89.0%로 세 분류 중 선별 안정성이 가장 높습니다.")

    add_heading(document, "6. 낮은 일치율의 원인")
    mismatch_rows = [
        ["날짜 범위 후보 미확보", metrics["candidate_missing_count"], "과거 원문은 DB에 있으나 과거 메일 날짜 규칙 밖에 발행된 기사입니다. AI가 볼 기회가 없으므로 수집·날짜 기준 문제로 분리합니다."],
        ["후보 중 AI·정책 제외", 80, "후보에는 있었지만 대표 기사 선정, 제외 기준, 섹션 자격 또는 메일 수량 정책으로 제외됐습니다."],
        ["후보에는 있으나 미선정", 27, "비교 가능한 후보였지만 최종 대표 기사로 선택되지 않았습니다. 우선순위 기준과 대표 기사 기준의 개선 대상입니다."],
    ]
    add_table(document, ["원인", "건수", "설명"], mismatch_rows, [2500, 900, 5960], font_size=8.2)
    add_callout(
        document,
        "해석",
        "전체 재현율 59.6%가 모두 AI 오류를 뜻하지 않습니다. 166건은 AI 입력 후보 자체에 없었고, AI가 실제로 비교할 수 있었던 509건 중에는 402건을 재현해 79.0%를 기록했습니다.",
        SOFT_GREEN,
    )

    add_heading(document, "7. 이번에 반영한 개선")
    improvements = [
        "과거 원본을 후보군에 강제로 넣는 검증 방식을 제거하고, DB에 실제 존재하는 후보만으로 다시 평가했습니다.",
        "AI가 선정한 순서를 서버 후처리에서 보존해 AI 판단이 점수 기반 재정렬로 덮이지 않도록 수정했습니다.",
        "동국대 관련성을 제목·요약뿐 아니라 발행인·기자명·주석·게시판·본문에서도 확인하도록 보강했습니다.",
        "빈 URL끼리 같은 기사로 오인하던 중복 판정을 수정했습니다.",
        "같은 학교명이 있다는 이유만으로 서로 다른 주제가 합쳐지지 않도록 백엔드 중복 기준을 엄격하게 조정하고, 의미 기반 주제 묶음은 AI가 담당하게 했습니다.",
        "정책 버전을 올려 과거 AI 캐시가 새 기준에 섞이지 않도록 했습니다.",
        "수집률, 후보 내 AI 재현율, 선택 기사 정확도, 전체 재현율을 분리해 낮은 수치의 원인을 바로 찾을 수 있도록 했습니다.",
    ]
    for item in improvements:
        add_bullet(document, item)

    add_heading(document, "8. 운영 권고")
    add_bullet(document, "수집 목표: 상위 분류별 날짜 범위 후보 확보율을 95% 이상으로 관리합니다. 특히 대학 [교육] 후보군의 검색어·매체 풀을 우선 확장합니다.")
    add_bullet(document, "AI 목표: 후보 내 AI 재현율을 85% 이상으로 올리고, 동국대 [법인/건학위]에서 제외된 80건과 미선정 27건을 월별 학습 자료로 사용합니다.")
    add_bullet(document, "대시보드: AI 인사이트에서 월별 기준 변경 문장, 변경 이유, 근거가 된 사용자 행동을 함께 표시하고 관리자가 불필요한 인사이트를 삭제할 수 있게 유지합니다.")
    add_bullet(document, "검증: URL 일치만 보지 않고 동일 주제 대표 기사까지 함께 검토하되, 세부 감사 자료에는 URL·제목·주제 일치를 각각 남깁니다.")

    landscape = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(landscape, landscape=True)
    add_heading(document, "부록 A. 날짜별 비교 결과")
    add_body(document, "각 날짜의 원본 수, 날짜 범위 후보, AI 선정 수와 세 가지 핵심 비율을 함께 표시합니다.")
    date_rows = []
    for mail_date, payload in sorted(comparison["dates"].items()):
        row = payload["summary"]
        date_rows.append([
            mail_date,
            row["홍보처 원본 기사"],
            row["날짜 범위 후보 확보"],
            row["AI 선정 기사"],
            row["URL 일치"],
            row["제목 일치"],
            row["동일 주제·대표 매체 차이"],
            row["총 일치"],
            f"{row['날짜 범위 후보율(%)']:.1f}%",
            f"{row['후보 내 AI 재현율(%)']:.1f}%",
            f"{row['주제 포함 일치율(%)']:.1f}%",
        ])
    add_table(
        document,
        ["날짜", "원본", "후보", "AI", "URL", "제목", "주제", "총 일치", "후보율", "AI 재현", "전체 재현"],
        date_rows,
        [1500, 900, 900, 780, 720, 720, 720, 900, 1260, 1440, 1560],
        total_width=LANDSCAPE_WIDTH_DXA,
        font_size=7.1,
    )

    portrait = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(portrait, landscape=False)
    add_heading(document, "부록 B. 대표 불일치 사례")
    examples = []
    for _, payload in sorted(comparison["dates"].items()):
        for item in payload["details"]:
            if item["비교 결과"] not in {"일치", "동일 주제 대표 매체 차이"}:
                examples.append(item)
    for item in examples[:15]:
        heading = document.add_paragraph()
        keep_paragraph_with_next(heading)
        first = heading.add_run(f"{item['메일 날짜']} · {item['비교 결과']}  ")
        set_run_font(first, 10, True, ACCENT)
        second = heading.add_run(item["원본 기사 제목"])
        set_run_font(second, 10, True, TEXT)
        add_body(document, item["차이/제외 이유"])

    add_callout(
        document,
        "최종 판단",
        (
            f"현재 시스템은 후보를 볼 수 있었던 기사에서는 {metrics['candidate_selection_recall_rate']:.1f}%를 재현하고 있습니다. "
            "따라서 다음 개선 우선순위는 ① 날짜 범위 후보 확보율 95% 달성, ② 동국대 [법인/건학위]의 우선순위·대표 기사 기준 보정, "
            "③ 월별 사용자 편집 로그를 이용한 기준표 개선입니다."
        ),
        SOFT_BLUE,
    )

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
